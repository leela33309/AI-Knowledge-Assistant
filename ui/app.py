import streamlit as st
import sys
import os
import time
import pandas as pd
import requests
from io import BytesIO
from docx import Document
from dotenv import load_dotenv

# =========================
# LOAD ENV
# =========================
load_dotenv()
APP_PASSWORD = os.getenv("APP_PASSWORD", "Leela@2002")

# =========================
# PATH FIX
# =========================
sys.path.append(
    os.path.abspath(
        os.path.join(
            os.path.dirname(__file__),
            ".."
        )
    )
)

# =========================
# IMPORT MODULES
# =========================
from core.file_loader import load_uploaded_file
from core.vector_store import create_index, search_index
from core.logger import log_query

# =========================
# STREAMLIT SETTINGS
# =========================
os.environ["STREAMLIT_WATCHER_TYPE"] = "none"

st.set_page_config(
    page_title="AI Knowledge Assistant",
    layout="wide"
)

st.title("🤖 AI Knowledge Assistant")

# =========================
# LOGIN
# =========================
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:

    st.subheader("🔐 Secure Login")

    password = st.text_input(
        "Enter Access Key",
        type="password"
    )

    if st.button("Login"):

        if password == APP_PASSWORD:
            st.session_state.authenticated = True
            st.success("Access Granted")
            st.rerun()

        else:
            st.error("Invalid Password")

    st.stop()

# =========================
# SESSION STATE
# =========================
if "messages" not in st.session_state:
    st.session_state.messages = []

if "input_text" not in st.session_state:
    st.session_state.input_text = ""

if "file_chunks" not in st.session_state:
    st.session_state.file_chunks = []

# =========================
# SIDEBAR
# =========================
with st.sidebar:

    st.header("⚙ Controls")

    if st.button("🗑 Clear Chat"):
        st.session_state.messages = []
        st.rerun()

    if st.button("🧹 Remove Uploaded File"):
        st.session_state.file_chunks = []
        st.success("File Removed")

    if st.button("🔄 Reindex File"):

        if st.session_state.file_chunks:
            create_index(
                st.session_state.file_chunks
            )
            st.success("Reindexed")

        else:
            st.warning("No File Uploaded")

# =========================
# FILE UPLOAD
# =========================
uploaded_file = st.file_uploader(
    "📄 Upload TXT / PDF / DOCX / CSV",
    type=["txt", "pdf", "docx", "csv"]
)

if uploaded_file:

    try:
        file_chunks = load_uploaded_file(
            uploaded_file
        )

        if file_chunks:

            st.session_state.file_chunks = file_chunks

            create_index(file_chunks)

            st.success(
                "✅ File Uploaded & Indexed"
            )

        else:
            st.warning("Could not read file")

    except Exception as e:
        st.error(str(e))

# =========================
# SMALL TALK
# =========================
def is_small_talk(text):

    greetings = [
        "hi",
        "hello",
        "hey",
        "good morning",
        "good evening",
        "how are you"
    ]

    return text.lower().strip() in greetings

# =========================
# WORD EXPORT
# =========================
def create_word_file(question, answer):

    doc = Document()

    doc.add_heading(
        "AI Assistant Response",
        0
    )

    doc.add_paragraph(
        f"Question: {question}"
    )

    doc.add_paragraph(
        f"Answer: {answer}"
    )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer

# =========================
# CSV EXPORT
# =========================
def create_csv_file(question, answer):

    df = pd.DataFrame({
        "Question": [question],
        "Answer": [answer]
    })

    return df.to_csv(
        index=False
    ).encode("utf-8")

# =========================
# MAIN QUERY FUNCTION
# =========================
def handle_query():

    try:

        query = st.session_state.input_text.strip()

        if not query:
            return

        start_time = time.time()

        st.session_state.messages.append({
            "role": "user",
            "content": query
        })

        sources = []

        # =====================
        # SMALL TALK
        # =====================
        if is_small_talk(query):

            answer = "Hi 😊 How can I help you today?"

        else:

            # =====================
            # API MODE
            # =====================
            try:

                response = requests.post(
                    "http://127.0.0.1:8000/ask",
                    json={
                        "question": query
                    }
                )

                data = response.json()

                answer = data.get(
                    "answer",
                    "No response generated."
                )

                api_sources = data.get(
                    "sources",
                    []
                )

                for item in api_sources:

                    if isinstance(item, dict):

                        src = item.get(
                            "source",
                            "Knowledge Base"
                        )

                        sources.append(src)

            except Exception as e:

                answer = f"⚠️ API Error: {str(e)}"

        # =====================
        # SOURCES
        # =====================
        if sources:

            answer += "\n\n📌 Sources:\n"
            answer += "\n".join(sources)

        # =====================
        # LOGGING
        # =====================
        latency = round(
            time.time() - start_time,
            2
        )

        try:
            log_query(
                query,
                answer,
                latency,
                len(sources)
            )
        except:
            pass

        st.session_state.messages.append({
            "role": "bot",
            "content": answer
        })

        st.session_state.input_text = ""

    except Exception as e:

        st.session_state.messages.append({
            "role": "bot",
            "content": f"⚠️ Error: {str(e)}"
        })

# =========================
# CHAT DISPLAY
# =========================
for i, msg in enumerate(
    st.session_state.messages
):

    if msg["role"] == "user":

        st.markdown(
            f"""
            <div style="text-align:right;
            margin:12px 0;">
                <div style="
                display:inline-block;
                background:#2b313e;
                color:white;
                padding:12px 16px;
                border-radius:18px;
                max-width:70%;
                text-align:left;">
                {msg["content"]}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    else:

        st.markdown(
            f"""
            <div style="text-align:left;
            margin:12px 0;">
                <div style="
                display:inline-block;
                background:#f1f3f6;
                color:black;
                padding:12px 16px;
                border-radius:18px;
                max-width:70%;
                text-align:left;
                white-space:pre-wrap;">
                🤖 {msg["content"]}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # LAST RESPONSE DOWNLOAD
        if i == len(
            st.session_state.messages
        ) - 1:

            if i > 0:

                question = st.session_state.messages[
                    i - 1
                ]["content"]

                col1, col2 = st.columns(2)

                with col1:

                    st.download_button(
                        "📄 Download Word",
                        data=create_word_file(
                            question,
                            msg["content"]
                        ),
                        file_name="AI_Response.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                with col2:

                    st.download_button(
                        "📊 Download CSV",
                        data=create_csv_file(
                            question,
                            msg["content"]
                        ),
                        file_name="AI_Response.csv",
                        mime="text/csv"
                    )

# =========================
# INPUT BOX
# =========================
st.text_input(
    "Type your question:",
    key="input_text",
    placeholder="Ask anything...",
    on_change=handle_query
)

# =========================
# CSS FIX
# =========================
st.markdown("""
<style>
div[data-baseweb="input"] input{
border:1px solid #cccccc !important;
padding:10px !important;
}
</style>
""", unsafe_allow_html=True)